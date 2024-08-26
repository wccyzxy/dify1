from docx import Document
from copy import deepcopy


def doc_to_json(doc_path):
    # ����Word�ĵ�
    doc = Document(doc_path)

    full_text = []

    body = doc.element.body
    table_title = ''

    # ����bodyԪ���е�������Ԫ��
    for element in body:
        if element.tag.endswith('p'):  # Check if it's a paragraph
            # ��ȡ�����ı�
            para_text = element.text
            if para_text:
                # ��������ʽ
                style_name = ''
                try:
                    style_name = doc.styles.element.style_lst[int(element.style)].name_val if element.style else ''
                except Exception as e:
                    style_name = ''
                if style_name == 'heading 1':
                    full_text.append({'content': para_text, 'type': 'text', 'level': 1})
                elif style_name == 'heading 2':
                    full_text.append({'content': para_text, 'type': 'text', 'level': 2})
                elif style_name == 'heading 3':
                    full_text.append({'content': para_text, 'type': 'text', 'level': 3})
                elif style_name == 'heading 4':
                    full_text.append({'content': para_text, 'type': 'text', 'level': 4})
                elif style_name == 'heading 5':
                    full_text.append({'content': para_text, 'type': 'text', 'level': 5})
                else:
                    if para_text.startswith('�� '):
                        table_title = para_text
                    else:
                        full_text.append({'content': para_text, 'type': 'text', 'level': -1})
        elif element.tag.endswith('tbl'):  # Check if it's a table
            # ��ȡ�������
            table_data = []
            for row in element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr'):
                row_data = []
                for cell in row.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc'):
                    texts = cell.xpath('descendant::text()')
                    cell_text = "".join([node.strip() for node in texts])
                    row_data.append(cell_text)
                table_data.append(row_data)
            full_text.append({'content': table_data, 'title': table_title, 'type': 'table', 'level': -1})

    return full_text


def doc_to_json_with_level(doc_path):
    # ����Word�ĵ�
    doc = Document(doc_path)

    doc_json = {
        'type': 'document',
        'level': -1,
        'content': '',
        'children': []
    }

    section_stack = [doc_json]
    table_title = ''

    body = doc.element.body

    # ����bodyԪ���е�������Ԫ��
    for element in body:
        if element.tag.endswith('p'):  # Check if it's a paragraph
            # ��ȡ�����ı�
            para_text = element.text
            if para_text:
                # ��������ʽ
                style_name = ''
                try:
                    style_name = doc.styles.element.style_lst[int(element.style)].name_val if element.style else ''
                except Exception as e:
                    style_name = ''
                level = -1
                if style_name == 'heading 1':
                    level = 1
                elif style_name == 'heading 2':
                    level = 2
                elif style_name == 'heading 3':
                    level = 3
                elif style_name == 'heading 4':
                    level = 4
                elif style_name == 'heading 5':
                    level = 5
                if level > 0:
                    # ����һ���µ�section
                    new_section = {
                        'type': 'text',
                        'level': level,
                        'content': para_text,
                        'children': []
                    }
                    # ����²㼶С�ڵ���ջ��Ԫ�صĲ㼶���򵯳�ջ��Ԫ��ֱ���ҵ����ʵĲ㼶
                    while section_stack and section_stack[-1]['level'] and level <= section_stack[-1]['level']:
                        section_stack.pop()

                    # ���½ڵ���ӵ���ǰ�㼶��ĩβ
                    section_stack[-1]['children'].append(new_section)
                    section_stack.append(new_section)
                else:
                    if para_text.startswith('�� '):
                        table_title = para_text
                    else:
                        current_section = section_stack[-1]
                        current_section['children'].append({
                            'type': 'text',
                            'level': -1,
                            'content': para_text
                        })
        elif element.tag.endswith('tbl'):  # Check if it's a table
            # ��ȡ�������
            table_data = []
            for row in element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr'):
                row_data = []
                for cell in row.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc'):
                    texts = cell.xpath('descendant::text()')
                    cell_text = "".join([node.strip() for node in texts])
                    row_data.append(cell_text)
                table_data.append(row_data)
            current_section = section_stack[-1]
            current_section['children'].append({
                'type': 'table',
                'title': table_title,
                'content': table_data,
                'level': -1
            })

    return doc_json


def save_content_to_md(content_list):
    content = ''
    for item in content_list:
        if item['type'] == 'text' and item['level'] > 0:
            content += '#' * item['level'] + ' ' + item['content'] + '\n\n'
        elif item['type'] == 'table':
            content += f"**Table Title:** {item.get('title', 'No Title')}\n\n"
            for row in item['content']:
                content += '|'.join(row) + '\n'
            content += '\n\n'
        elif item['type'] == 'text':
            content += item['content'] + '\n\n'
    return content


def copy_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    new_tbl = deepcopy(tbl)
    p.addnext(new_tbl)


def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def move_paragraph_after(paragraph, paragraph_after):
    p, p_after = paragraph._p, paragraph_after._p
    p_after.addnext(p)


def delete_table_with_title(document, expect_text):
    allTables = document.tables
    for activeTable in allTables:
        if activeTable.cell(0, 0).paragraphs[0].text == expect_text:
            activeTable._element.getparent().remove(activeTable._element)


def extract_tables_from_docx(file_path):
    # ���ĵ�
    doc = Document(file_path)

    result = ""
    # �����ĵ��е����б��
    for i, table in enumerate(doc.tables):
        result += f"Table {i + 1}:\n"
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text)
            result += " | ".join(row_data) + "\n"
        result += "\n"

    return result


def extract_paragraphs_from_docx(file_path):
    # ��docx�ļ�
    doc = Document(file_path)

    # ��ʼ��һ�����ַ������洢�ı�
    full_text = []

    # �����ĵ��е����ж���
    for para in doc.paragraphs:
        full_text.append(para.text)

    # �����ж�����ı����ӳ�һ���ַ���
    return '\n'.join(full_text)


def extract_text_from_docx(file_path):
    # ��docx�ļ�
    doc = Document(file_path)

    # ��ʼ��һ�����б����洢�ı�
    full_text = []

    # �����ĵ��е����ж���
    for para in doc.paragraphs:
        text = para.text.strip()  # ȥ�������ı���ǰ��հ�
        if text:  # ��������ı���Ϊ�գ�����ӵ��б���
            full_text.append(text)

    # �����ĵ��е����б��
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()  # ȥ�������ı���ǰ��հ�
                    if text:  # ��������ı���Ϊ�գ�����ӵ��б���
                        full_text.append(text)

    # �����ж���ͱ����ı����ӳ�һ���ַ���
    return '\n'.join(full_text)

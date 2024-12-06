import json
import os
import sys
import re
from docx.oxml.ns import qn
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from lxml import etree
from core.rag.extractor.util.html_table_to_docx import html_table_to_docx


class DocxUtils:
    def __init__(self, chunk_size=1500):
        self.doc = None
        self.chunk_size = chunk_size

    ############################################################
    # 提取docx中的文本、表格、段落样式
    ############################################################
    def doc_to_json(self, doc_path):
        # 加载Word文档
        doc = Document(doc_path)

        full_text = []

        body = doc.element.body

        # 遍历body元素中的所有子元素
        for element in body:
            if element.tag.endswith('p'):  # Check if it's a paragraph
                # 提取段落文本
                para_text = element.text
                if para_text:
                    # 检查段落样式
                    style_name = ''
                    try:
                        style_name = doc.styles.element.style_lst[int(element.style)].name_val if element.style else ''
                    except Exception as e:
                        style_name = ''
                    if style_name == 'heading 1':
                        full_text.append({'content': para_text, 'type': 'text', 'level': 1})
                    elif style_name == 'heading 2':
                        full_text.append({'content': para_text, 'type': 'text', 'level': 2})
                    elif style_name == 'heading 3':
                        full_text.append({'content': para_text, 'type': 'text', 'level': 3})
                    elif style_name == 'heading 4':
                        full_text.append({'content': para_text, 'type': 'text', 'level': 4})
                    elif style_name == 'heading 5':
                        full_text.append({'content': para_text, 'type': 'text', 'level': 5})
                    else:
                        full_text.append({'content': para_text, 'type': 'text', 'level': -1})
            elif element.tag.endswith('tbl'):  # Check if it's a table
                # 提取表格数据
                table_data = []
                for row in element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr'):
                    row_data = []
                    for cell in row.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc'):
                        texts = cell.xpath(
                            'descendant::w:t[not(ancestor::w:commentRangeStart) and not(ancestor::w:commentRangeEnd)]')
                        cell_text = "".join([node.text.strip() if node.text else "" for node in texts])
                        row_data.append(cell_text)
                    table_data.append(row_data)
                full_text.append({'content': table_data, 'type': 'table', 'level': -1})

        full_text

    def extract_paragraph_content(self, paragraph):
        content = []
        # 遍历所有子元素,按顺序处理文本和公式
        for element in paragraph:
            if element.tag.endswith('}r'):  # 处理普通文本
                is_subscript = False
                is_superscript = False
                e_rPr = element.rPr
                if e_rPr is not None:
                    is_subscript = e_rPr.subscript
                    is_superscript = e_rPr.superscript

                # 提取文本，使用正确的命名空间
                text = element.text

                if text:
                    if is_subscript:
                        content.append(f"$_{text}$")  # 使用 _{text} 格式表示下标
                    elif is_superscript:
                        content.append(f"$^{text}$")
                    else:
                        content.append(text)
            elif element.tag.endswith('}oMath'):  # 处理公式
                markdown_equation = self.omml_to_markdown(element)
                content.append(markdown_equation)

        # 将所有内容按顺序拼接
        return "".join(content)

    def extract_table_data(self, table):
        """提取表格数据"""
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_data.append(cell_text)
            table_data.append(row_data)
        return table_data

    def extract_section_content(self, section_part):
        """处理页眉或页脚的内容"""
        content = []

        # 处理段落
        for paragraph in section_part.paragraphs:
            text = self.extract_paragraph_content(paragraph.part.element)
            if text.strip():
                content.append({
                    'type': 'text',
                    'level': -1,
                    'content': text
                })

        # 处理表格
        for table in section_part.tables:
            table_data = self.extract_table_data(table)
            if table_data:
                content.append({
                    'type': 'table',
                    'level': -1,
                    'content': table_data
                })

        return content

    # 将OMML转换为简单的Markdown公式
    def omml_to_markdown(self, omath):
        equation = etree.tostring(omath, encoding='unicode')
        equation = self.handle_special_characters(equation)
        # 替换一些常见的OMML标签为Markdown语法
        equation = re.sub(r'<m:oMath.*?>(.*?)</m:oMath>', r'$\1$', equation, flags=re.DOTALL)
        equation = re.sub(r'<m:f>(.*?)</m:f>', r'\\frac{\1}', equation, flags=re.DOTALL)
        equation = re.sub(r'<m:num>(.*?)</m:num>', r'{\1}', equation, flags=re.DOTALL)
        equation = re.sub(r'<m:den>(.*?)</m:den>', r'{\1}', equation, flags=re.DOTALL)
        equation = re.sub(r'<m:sup>(.*?)</m:sup>', r'^{\1}', equation, flags=re.DOTALL)
        equation = re.sub(r'<m:sub>(.*?)</m:sub>', r'_{\1}', equation, flags=re.DOTALL)
        equation = re.sub(r'<m:r>(.*?)</m:r>', r'\1', equation, flags=re.DOTALL)
        equation = re.sub(r'<m:t>(.*?)</m:t>', r'\1', equation, flags=re.DOTALL)

        # 移除所有剩余的XML标签
        equation = re.sub(r'<.*?>', '', equation)

        # 清理空白字符
        equation = re.sub(r'\s+', ' ', equation).strip()

        return equation

    @staticmethod
    def handle_special_characters(equation):
        # 定义特殊字符映射
        symbol_map = {
            'F065': '\\Alpha',  # 大写 Alpha
            'F066': '\\Beta',  # 大写 Beta
            'F067': '\\Chi',  # 大写 Chi
            # ... 添加更多映射
        }

        # 使用正则表达式查找并替换特殊字符
        def replace_symbol(match):
            char_code = match.group(1)
            return symbol_map.get(char_code, f'\\symbol{{{char_code}}}')

        equation = re.sub(r'<w:sym\s+w:font="Symbol"\s+w:char="(F[0-9A-F]{3})"/>', replace_symbol, equation)

        return equation

    @staticmethod
    def clean_string(s):
        # 将中文空格替换为英文空格
        s = s.replace('\u3000', ' ')
        # 去除前后的空白字符
        s = s.strip()
        # 使用正则表达式替换连续的空格为单个空格
        s = re.sub(r'\s+', ' ', s)
        return s

    def merge_deal_data(self, existing_deals, new_deals):
        """合并相同type的deal数据"""
        # 用字典存储现有的deals,以type为key
        deal_map = {}
        for deal in existing_deals:
            deal_map[deal['type']] = deal

        # 合并新的deals
        for new_deal in new_deals:
            if new_deal['type'] in deal_map:
                # 如果已存在相同type,合并items
                for item in new_deal['items']:
                    if item in deal_map[new_deal['type']]['items']:
                        continue
                    deal_map[new_deal['type']]['items'].append(item)
            else:
                # 如果是新type,直接添加
                deal_map[new_deal['type']] = new_deal

        return list(deal_map.values())

    def get_style_by_id(self, doc, style_id):
        try:
            # 获取文档中的所有样式
            styles = doc.styles.element.style_lst

            # 遍历样式列表查找匹配的style_id
            for style in styles:
                if style.styleId == style_id:
                    return style

        except Exception as e:
            print(f"获取样式时发生错误: {str(e)}")
            return None

    # 根据预处理json，将预处理模板转换为最终模板，生成最终模板json
    def doc_to_json_with_level(self, pre_deal_doc_path, pre_deal_json_path, final_template_path, final_template_name):
        # 创建html
        html_path = pre_deal_doc_path.replace('.docx', '.html')
        self.docx_to_html(pre_deal_doc_path, html_path)

        # 加载Word文档
        doc = Document(pre_deal_doc_path)

        pre_deal_json_data = {}
        with open(pre_deal_json_path, 'r', encoding='utf-8') as f:
            pre_deal_json_data = json.load(f)

        final_template_json_data = {
            'type': 'document',
            'title': final_template_name,
            'deal': [],
            'header': [],
            'header_deal': [],
            'level': -1,
            'children': [],
            'attachments': [],
            'path': final_template_path.split('/')[-1].replace('.final', ''),
            'finished': 0
        }

        # 处理页面页脚
        for section in doc.sections:
            header_content = self.extract_section_content(section.header)
            if header_content:
                final_template_json_data['header'].append({
                    'type': 'header',
                    'level': -1,
                    'children': header_content
                })
                header_text = ''
                for content in header_content:
                    if content['type'] == 'text':
                        header_text += content['content']
                    elif content['type'] == 'table':
                        for row in content['content']:
                            header_text += '|'.join(row) + '\n'
                deal_data = self.get_deal_data(doc, header_text, pre_deal_json_data, html_path)
                if len(deal_data) > 0:
                    final_template_json_data['header_deal'] = self.merge_deal_data(
                        final_template_json_data['header_deal'], deal_data)

        section_stack = [final_template_json_data]
        table_title = ''

        body = doc.element.body

        # 遍历body元素中的所有子元素
        heading_pattern = re.compile(r'^(\d+(\.\d+)*[.)]?)')
        for element in body:
            if element.tag.endswith('p'):  # Check if it's a paragraph
                # 提取段落文本
                para_text = self.extract_paragraph_content(element)

                # 去掉目录
                if para_text.replace(' ', '') == "目录":
                    continue

                if para_text:
                    style_name = ''
                    try:
                        style = self.get_style_by_id(doc, element.style)
                        style_name = style.name_val if style is not None else element.style
                        if style_name.startswith("toc") or style_name.startswith("TOC"):
                            continue
                        if style_name == '' and element.style:
                            style_name = element.style
                    except Exception as e:
                        style_name = element.style if element.style else ''
                        if element.style and element.style.startswith("TOC"):
                            continue

                    level = -1
                    if style_name == 'heading 1' or style_name == 'Heading1':
                        level = 1
                    elif style_name == 'heading 2' or style_name == 'Heading2':
                        level = 2
                    elif style_name == 'heading 3' or style_name == 'Heading3':
                        level = 3
                    elif style_name == 'heading 4' or style_name == 'Heading4':
                        level = 4
                    elif style_name == 'heading 5' or style_name == 'Heading5':
                        level = 5
                    elif style_name == 'heading 6' or style_name == 'Heading6':
                        level = 6
                    else:
                        match = heading_pattern.match(para_text)
                        if match:
                            # 计算标题序号中的点的数量来确定标题的层级
                            level = len(match.group(1).split('.'))
                            # 如果层级大于3，则从第4级开始依次增加
                            if level > 3:
                                level += 6
                    if level > 0:
                        # 创建一个新的section
                        new_section = {
                            'type': 'text',
                            'level': level,
                            'content': para_text,
                            'children': [],
                            'deal': [],
                            'finished': 0
                        }
                        deal_data = self.get_deal_data(doc, para_text, pre_deal_json_data, html_path)
                        if len(deal_data) > 0:
                            new_section['deal'] = deal_data
                        # 如果新层级小于等于栈顶元素的层级，则弹出栈顶元素直到找到合适的层级
                        while section_stack and section_stack[-1]['level'] and level <= section_stack[-1]['level']:
                            section_stack.pop()

                        # 将新节点添加到当前层级的末尾
                        section_stack[-1]['children'].append(new_section)
                        section_stack.append(new_section)
                    else:
                        if para_text.startswith('表 ') or para_text.endswith('表') or para_text.startswith('表'):
                            table_title = para_text
                        else:
                            current_section = section_stack[-1]
                            current_section['children'].append({
                                'type': 'text',
                                'level': -1,
                                'content': para_text
                            })
                            deal_data = self.get_deal_data(doc, para_text, pre_deal_json_data, html_path)
                            if len(deal_data) > 0:
                                current_section['deal'] = self.merge_deal_data(current_section['deal'], deal_data)
            elif element.tag.endswith('tbl'):  # Check if it's a table
                # 提取表格数据
                table_data = []
                table_data_text = ''
                for row in element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr'):
                    row_data = []
                    for cell in row.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc'):
                        texts = cell.xpath(
                            'descendant::w:t[not(ancestor::w:commentRangeStart) and not(ancestor::w:commentRangeEnd)]')
                        cell_text = "".join([node.text.strip() if node.text else "" for node in texts])
                        row_data.append(cell_text)
                    table_data.append(row_data)
                    table_data_text += '|'.join(row_data) + '\n'
                current_section = section_stack[-1]
                current_section['children'].append({
                    'type': 'table',
                    'title': table_title,
                    'content': table_data,
                    'level': -1
                })
                deal_data = self.get_deal_data(doc, table_title + "\n" + table_data_text, pre_deal_json_data, html_path)
                if len(deal_data) > 0:
                    current_section['deal'] = self.merge_deal_data(current_section['deal'], deal_data)

                table_title = ''

        doc.save(final_template_path)
        return final_template_json_data

    def doc_to_json_with_level(self, docx_path):
        # 加载Word文档
        doc = Document(docx_path)

        final_template_json_data = {
            'type': 'document',
            'deal': [],
            'header': [],
            'header_deal': [],
            'level': -1,
            'children': [],
            'attachments': [],
            'finished': 0
        }

        # 处理页面页脚
        for section in doc.sections:
            header_content = self.extract_section_content(section.header)
            if header_content:
                final_template_json_data['header'].append({
                    'type': 'header',
                    'level': -1,
                    'children': header_content
                })
                header_text = ''
                for content in header_content:
                    if content['type'] == 'text':
                        header_text += content['content']
                    elif content['type'] == 'table':
                        for row in content['content']:
                            header_text += '|'.join(row) + '\n'

        section_stack = [final_template_json_data]
        table_title = ''

        body = doc.element.body

        # 遍历body元素中的所有子元素
        heading_pattern = re.compile(r'^(\d+(\.\d+)*[.)]?)')
        for element in body:
            if element.tag.endswith('p'):  # Check if it's a paragraph
                # 提取段落文本
                para_text = self.extract_paragraph_content(element)

                # 去掉目录
                if para_text.replace(' ', '') == "目录":
                    continue

                if para_text:
                    style_name = ''
                    try:
                        style = self.get_style_by_id(doc, element.style)
                        style_name = style.name_val if style is not None else element.style
                        if style_name.startswith("toc") or style_name.startswith("TOC"):
                            continue
                        if style_name == '' and element.style:
                            style_name = element.style
                    except Exception as e:
                        style_name = element.style if element.style else ''
                        if element.style and element.style.startswith("TOC"):
                            continue

                    level = -1
                    if style_name == 'heading 1' or style_name == 'Heading1':
                        level = 1
                    elif style_name == 'heading 2' or style_name == 'Heading2':
                        level = 2
                    elif style_name == 'heading 3' or style_name == 'Heading3':
                        level = 3
                    elif style_name == 'heading 4' or style_name == 'Heading4':
                        level = 4
                    elif style_name == 'heading 5' or style_name == 'Heading5':
                        level = 5
                    elif style_name == 'heading 6' or style_name == 'Heading6':
                        level = 6
                    else:
                        match = heading_pattern.match(para_text)
                        if match:
                            # 计算标题序号中的点的数量来确定标题的层级
                            length = len(match.group(1).split('.'))
                            if length > 3:
                                level = length + 6
                    if level > 0:
                        # 创建一个新的section
                        new_section = {
                            'type': 'text',
                            'level': level,
                            'content': para_text,
                            'children': [],
                            'deal': [],
                            'finished': 0
                        }
                        # 如果新层级小于等于栈顶元素的层级，则弹出栈顶元素直到找到合适的层级
                        while section_stack and section_stack[-1]['level'] and level <= section_stack[-1]['level']:
                            section_stack.pop()

                        # 将新节点添加到当前层级的末尾
                        section_stack[-1]['children'].append(new_section)
                        section_stack.append(new_section)
                    else:
                        if para_text.startswith('表 ') or para_text.endswith('表') or para_text.startswith(
                                '表') or para_text == "方案摘要":
                            table_title = para_text
                        else:
                            current_section = section_stack[-1]
                            current_section['children'].append({
                                'type': 'text',
                                'level': -1,
                                'content': para_text
                            })
            elif element.tag.endswith('tbl'):  # Check if it's a table
                # 提取表格数据
                table_data = []
                table_data_text = ''
                for row in element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr'):
                    row_data = []
                    for cell in row.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc'):
                        texts = cell.xpath(
                            'descendant::w:t[not(ancestor::w:commentRangeStart) and not(ancestor::w:commentRangeEnd)]')
                        cell_text = "".join([node.text.strip() if node.text else "" for node in texts])
                        row_data.append(cell_text)
                    table_data.append(row_data)
                    table_data_text += '|'.join(row_data) + '\n'
                current_section = section_stack[-1]
                current_section['children'].append({
                    'type': 'table',
                    'title': table_title,
                    'content': table_data,
                    'level': -1
                })
                table_title = ''
        return final_template_json_data

    def split_by_title(self, content, result, parent_titles=[]):
        if content['type'] == 'table':
            title = self.clean_string(content['title'])
            full_title = '\n'.join([t for t, level in parent_titles] + [title]) + "\n"
            for row in content['content']:
                row_text = self.clean_string('|'.join(row)) + "\n"
                if result[-1].get("content", "") in full_title and len(parent_titles) > 0:
                    result[-1] = {"content": full_title, "metadata": {}}
                elif full_title not in result[-1].get("content", "") and len(parent_titles) > 0:
                    result.append({"content": full_title + row_text, "metadata": {}})
                elif len(result[-1].get("content", "")) + len(row_text) > self.chunk_size:
                    result.append({"content": full_title + row_text, "metadata": {}})
                else:
                    result[-1] = {"content": result[-1].get("content", "") + row_text, "metadata": {}}
        elif content['level'] > 0:
            text = self.clean_string(content.get('content', ''))
            relevant_parents = [(title, level) for title, level in parent_titles if level < content['level']]
            parent_titles = relevant_parents + [(text, content['level'])]
        else:
            full_title = '\n'.join([t for t, level in parent_titles]) + '\n'
            text = self.clean_string(content.get('content', '')) + "\n"
            if len(parent_titles) > 0:
                if result[-1].get("content", "") in full_title:
                    result[-1] = {"content": full_title + text, "metadata": {}}
                elif full_title not in result[-1].get("content", ""):
                    result.append({"content": full_title + text, "metadata": {}})
                else:
                    if len(result[-1]) + len(text) > self.chunk_size:
                        result.append({"content": full_title + text, "metadata": {}})
                    else:
                        result[-1] = {"content": result[-1].get("content", "") + text, "metadata": {}}
            else:
                if len(result[-1]) + len(text) > self.chunk_size:
                    result.append({"content": text, "metadata": {}})
                else:
                    result[-1] = {"content": result[-1].get("content", "") + text, "metadata": {}}

        if 'children' in content.keys() and len(content['children']) > 0:
            for child in content['children']:
                result = self.split_by_title(child, result, parent_titles)
        else:
            text = "\n".join([t for t, level in parent_titles]) + "\n"
            if len(parent_titles) > 0:
                if result[-1].get("content", "") in text:
                    result[-1] = {"content": text, "metadata": {}}
                elif text not in result[-1].get("content", ""):
                    result.append({"content": text, "metadata": {}})

        return result

    def extract_docx_to_list(self, doc_path):
        doc_json = self.doc_to_json_with_level(doc_path)
        result = self.split_by_title(doc_json, [{"content": "", "metadata": {}}])
        return result

    ############################################################
    # 处理预处理json数据方法
    ############################################################
    def has_type(self, content, type):
        pattern = r'\{' + type + r'_\d+\}'
        return bool(re.search(pattern, content))

    def get_type_content(self, content, type):
        pattern = r'\{' + type + r'_\d+\}'
        matches = re.findall(pattern, content)
        return list(set(matches)) if matches else []

    def get_pre_deal_data(self, pre_deal_json_data, type_name, key):
        type_data = pre_deal_json_data.get(type_name, [])
        if len(type_data) > 0:
            for data in type_data:
                if data.get('key') == key:
                    return data
        return None

    def get_deal_data(self, doc, content, pre_deal_json_data, html_path):
        deal_data = []
        types = [
            'fill_blank',
            'alert_to_fill',
            'pic_template_to_write',
            'write_online',
            'suolueyu',
            'sort_file',
            'pick_subject',
            'table_extract',
            'copy'
        ]
        for type_name in types:
            if self.has_type(content, type_name):
                type_contents = self.get_type_content(content, type_name)
                if len(type_contents) > 0:
                    if type_name == 'fill_blank':
                        data = {
                            'type': type_name,
                            'items': [],
                            'dealed': 0
                        }
                        for content_item in type_contents:
                            key = content_item.replace('{', '').replace('}', '')
                            pre_deal_data = self.get_pre_deal_data(pre_deal_json_data, 'fill_blank', key)
                            if pre_deal_data:
                                data['items'].append({
                                    'key': content_item,
                                    'title': pre_deal_data['title'],
                                    'value': '',
                                    'prompt': pre_deal_data['prompt'],
                                    'dealed': 0,
                                    'prompts': []
                                })
                        if len(data['items']) > 0:
                            deal_data.append(data)
                    elif type_name == 'alert_to_fill':
                        data = {
                            'type': type_name,
                            'items': [],
                            'dealed': 0
                        }
                        for content_item in type_contents:
                            key = content_item.replace('{', '').replace('}', '')
                            pre_deal_data = self.get_pre_deal_data(pre_deal_json_data, 'alert_to_fill', key)
                            if pre_deal_data:
                                result_key = content.split(content_item)[1]
                                data['items'].append({
                                    'key': result_key,
                                    'value': '',
                                    'dealed': 0
                                })
                                self.replace_word(doc, content_item + result_key + content_item, result_key)
                                self.replace_table_word(doc, content_item + result_key + content_item, result_key)
                        if len(data['items']) > 0:
                            deal_data.append(data)
                    elif type_name == 'pic_template_to_write':
                        data = {
                            'type': type_name,
                            'items': [],
                            'key': '',
                            'value': '',
                            'dealed': 0
                        }
                        for content_item in type_contents:
                            key = content_item.replace('{', '').replace('}', '')
                            pre_deal_data = self.get_pre_deal_data(pre_deal_json_data, 'pic_template_to_write', key)
                            if pre_deal_data:
                                data['key'] = content_item
                                for item in pre_deal_data['templates']:
                                    data['items'].append({
                                        'title': item['title'],
                                        'content': item['content'],
                                        'prompt': item['prompt'],
                                        'prompts': []
                                    })
                        if len(data['items']) > 0:
                            deal_data.append(data)
                    elif type_name == 'write_online':
                        data = {
                            'type': type_name,
                            'items': [],
                            'dealed': 0
                        }
                        for content_item in type_contents:
                            key = content_item.replace('{', '').replace('}', '')
                            pre_deal_data = self.get_pre_deal_data(pre_deal_json_data, 'write_online', key)
                            if pre_deal_data:
                                data['items'].append({
                                    'key': content_item,
                                    'value': '',
                                    'prompt': pre_deal_data['prompt'],
                                    'dealed': 0,
                                    'prompts': []
                                })
                        if len(data['items']) > 0:
                            deal_data.append(data)
                    elif type_name == 'suolueyu':
                        pass
                    elif type_name == 'sort_file':
                        pass
                    elif type_name == 'pick_subject':
                        for content_item in type_contents:
                            data = {
                                'type': type_name,
                                'items': [],
                                'dealed': 0
                            }
                            key = content_item.replace('{', '').replace('}', '')
                            pre_deal_data = self.get_pre_deal_data(pre_deal_json_data, 'pick_subject', key)
                            if pre_deal_data:
                                for subject in pre_deal_data['subjects']:
                                    data['items'].append({
                                        'item': subject,
                                        'picked': 0,
                                    })
                            if len(data['items']) > 0:
                                deal_data.append(data)
                    elif type_name == 'table_extract':
                        for content_item in type_contents:
                            data = {
                                'type': type_name,
                                'items': [],
                                'dealed': 0
                            }
                            key = content_item.replace('{', '').replace('}', '')
                            pre_deal_data = self.get_pre_deal_data(pre_deal_json_data, 'table_extract', key)
                            if pre_deal_data:
                                table_name = content.split(content_item)[1]
                                table_origin_html = self.extract_table_from_html(
                                    f"{content_item}{table_name}{content_item}", html_path)
                                simple_table_html = self.simple_table_html(table_origin_html)
                                data['items'].append({
                                    'key': content_item,
                                    'title': table_name,
                                    'table_origin_html': table_origin_html,
                                    'table_simple_html': simple_table_html,
                                    'table_final_html': '',
                                    'prompt': pre_deal_data['prompt'],
                                    'dealed': 0
                                })
                            if len(data['items']) > 0:
                                deal_data.append(data)
                    elif type_name == 'copy':
                        for content_item in type_contents:
                            data = {
                                'type': type_name,
                                'items': [],
                                'dealed': 0
                            }
                            key = content_item.replace('{', '').replace('}', '')
                            pre_deal_data = self.get_pre_deal_data(pre_deal_json_data, 'copy', key)
                            if pre_deal_data:
                                data['items'].append({
                                    'key': content_item,
                                    'prompt': pre_deal_data['prompt'],
                                    'dealed': 0
                                })
                                deal_data.append(data)
                            if len(data['items']) > 0:
                                deal_data.append(data)
        return deal_data

    def docx_to_html(self, doc_path, html_path):
        import platform
        system = platform.system()

        if system == 'Windows':
            # Windows 环境使用 spire.doc
            from spire.doc import Document as sDocument, FileFormat
            document = sDocument()
            document.LoadFromFile(doc_path)
            document.SaveToFile(html_path, FileFormat.Html)
            document.Close()
        else:
            # Linux/Mac 环境使用 LibreOffice
            import subprocess
            import os

            # 确保输出目录存在
            os.makedirs(os.path.dirname(html_path), exist_ok=True)

            # 使用 LibreOffice 转换
            cmd = [
                'libreoffice',
                '--headless',
                '--convert-to',
                'html',
                '--outdir',
                os.path.dirname(html_path),
                doc_path
            ]
            subprocess.run(cmd, check=True)

            # LibreOffice 生成的文件名可能与期望的不同，需要重命名
            generated_html = doc_path.replace('.docx', '.html')
            if os.path.exists(generated_html) and generated_html != html_path:
                os.rename(generated_html, html_path)

    def extract_table_from_html(self, table_name, html_path):
        try:
            from bs4 import BeautifulSoup
            import re

            # 获取CSS文件路径
            css_path = html_path.replace('.html', '_style.css')

            # 读取HTML文件
            with open(html_path, 'r', encoding='utf-8') as f:
                html_content = f.read()

            # 读取CSS文件
            css_content = ''
            if os.path.exists(css_path):
                with open(css_path, 'r', encoding='utf-8') as f:
                    css_content = f.read()

            # 使用BeautifulSoup解析HTML
            soup = BeautifulSoup(html_content, 'html.parser')

            # 查找所有段落和表格
            elements = soup.find_all(['p', 'table'])

            # 遍历元素寻找目标表格
            for i, element in enumerate(elements):
                # 检查是否是段落且包含表名
                if (element.name == 'p' and
                        table_name in element.get_text().strip()):
                    # 查找下一个表格元素
                    for next_elem in elements[i + 1:]:
                        if next_elem.name == 'table':
                            # 提取表格HTML和相关样式
                            table_html = str(next_elem)
                            # 返回完整的HTML
                            return table_html.strip()
                    break

            # 如果没有找到表格，返回空字符串
            return ''

        except Exception as e:
            print(f"提取表格时发生错误: {str(e)}")
            return ''

    def _filter_css_for_table(self, css_content: str, table_classes: list) -> str:
        """
        从CSS内容中提取表格相关的样式

        Args:
            css_content: CSS文件内容
            table_classes: 表格使用的类名列表

        Returns:
            str: 过滤后的CSS内容
        """
        import re

        # 保存相关的CSS规则
        filtered_rules = []

        # 基本的表格选择器
        basic_selectors = ['table', 'th', 'td', 'tr']

        # 分割CSS规则
        rules = re.split(r'}\s*', css_content)

        for rule in rules:
            rule = rule.strip()
            if not rule:
                continue

            # 检查是否包含表格相关选择器
            selectors, styles = rule.split('{', 1)
            selectors = selectors.strip()

            # 检查是否是表格相关样式
            is_table_related = False

            # 检查基本选择器
            for selector in basic_selectors:
                if selector in selectors:
                    is_table_related = True
                    break

            # 检查类名
            for class_name in table_classes:
                if f'.{class_name}' in selectors:
                    is_table_related = True
                    break

            if is_table_related:
                filtered_rules.append(f'{selectors} {{{styles}}}')

        return '\n'.join(filtered_rules)

    def _preserve_style(self, old_element, new_element, style_attrs):
        """保留指定的样式属性"""
        # 处理直接的style属性
        preserved_styles = []
        if 'style' in old_element.attrs:
            old_styles = dict(style.strip().split(':') for style in old_element['style'].split(';') if ':' in style)
            for attr in style_attrs:
                if attr in old_styles:
                    preserved_styles.append(f"{attr}:{old_styles[attr]}")
        if 'valign' in old_element.attrs:
            preserved_styles.append(f"vertical-align:{old_element['valign']}")
        if preserved_styles:
            new_element['style'] = ';'.join(preserved_styles)

    def _process_nested_tags(self, parent_tag, new_parent, soup, text_styles):
        """递归处理嵌套的标签"""
        for child in parent_tag.children:
            if child.name:  # 如果是标签
                new_tag = soup.new_tag(child.name)
                # 保留样式
                self._preserve_style(child, new_tag, text_styles)

                # 递归处理子标签
                self._process_nested_tags(child, new_tag, soup, text_styles)
                new_parent.append(new_tag)
            elif child.string and child.string.strip():  # 如果是文本节点
                new_parent.append(child.string.strip())

    def simple_table_html(self, html_content):
        import platform
        system = platform.system()

        if system == 'Windows':
            return self.simple_table_html_win(html_content)
        return self.simple_table_html_linux(html_content)

    def simple_table_html_win(self, html_content):
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(html_content, 'html.parser')
        table_body = soup.find('table')
        rows = table_body.find_all('tr')

        formatted_soup = BeautifulSoup('''
        <table></table>
        '''.replace("\n", "").replace(" ", ""), 'html.parser')
        formatted_table = formatted_soup.table
        important_table_styles = ['width', 'table-layout', 'border-collapse', 'total_width']
        important_text_styles = ['font-weight', 'vertical-align', 'text-indent', 'text-align', 'margin-top']
        important_cell_styles = ['width', 'height', 'vertical-align', 'text-align', 'background-color']
        self._preserve_style(table_body, formatted_table, important_table_styles)
        for row in rows:
            cells = row.find_all('td')
            new_row = formatted_soup.new_tag('tr')
            self._preserve_style(row, new_row, important_cell_styles)
            for cell in cells:
                new_cell = formatted_soup.new_tag('td')
                content_tags = cell.find_all('p', recursive=False)
                if content_tags:
                    for p_tag in content_tags:
                        new_p = formatted_soup.new_tag('p')
                        self._preserve_style(p_tag, new_p, important_text_styles)

                        # 递归处理p标签内的所有嵌套内容
                        self._process_nested_tags(p_tag, new_p, formatted_soup, important_text_styles)
                        new_cell.append(new_p)
                else:
                    # 如果没有p标签，直接处理文本内容
                    cell_text = ''.join(cell.stripped_strings)
                    new_cell.string = cell_text

                # 保留colspan和rowspan属性
                for attr in ['colspan', 'rowspan']:
                    if attr in cell.attrs:
                        new_cell[attr] = cell[attr]
                self._preserve_style(cell, new_cell, important_cell_styles)
                new_row.append(new_cell)
            formatted_table.append(new_row)
        return formatted_soup.prettify().replace("\n", "").replace("  ", " ")

    def merge_styles(self, existing_style, new_styles):
        """合并样式字符串"""
        style_dict = {}

        # 解析现有样式
        if existing_style:
            existing_style_dict = self.parse_style(existing_style)
            style_dict = self.merge_style_dicts(style_dict, existing_style_dict)

        # 添加新样式
        for key, value in new_styles.items():
            style_dict[key] = value

        # 重新组合样式字符串
        return ';'.join([f'{k}:{v}' for k, v in style_dict.items()])

    @staticmethod
    def merge_style_dicts(style1, style2):
        """合并两个样式字典"""
        result = style1.copy()
        result.update(style2)
        return result

    @staticmethod
    def parse_style(style_str):
        """解析style字符串为字典"""
        style_dict = {}
        if style_str:
            for style_item in style_str.split(';'):
                if ':' in style_item:
                    key, value = style_item.split(':', 1)
                    style_dict[key.strip()] = value.strip()
        return style_dict

    def simple_table_html_linux(self, html_content):
        from bs4 import BeautifulSoup

        # 1. 移除多余的空tr标签
        html_content = re.sub(r'<tr>\s*</tr>', '', html_content)

        def convert_to_win_format(html_content):
            soup = BeautifulSoup(html_content, 'html.parser')

            # 1. 处理表格属性
            for table in soup.find_all('table'):
                attrs_to_style = ['width', 'height', 'align', 'valign']
                style_dict = {}
                for attr in attrs_to_style:
                    if attr in table.attrs:
                        style_dict[attr] = table[attr]
                        del table[attr]
                cols = soup.find_all('col')
                if cols:
                    total_width = 0
                    for col in cols:
                        width = float(col.get('width').replace('*', ''))
                        total_width += width
                        col.decompose()
                    if total_width > 0:
                        style_dict['total_width'] = total_width
                if style_dict:
                    existing_style = table.get('style', '')
                    existing_style_dict = self.parse_style(existing_style)
                    style_dict = self.merge_style_dicts(existing_style_dict, style_dict)
                    style_str = ';'.join([f'{k}:{v}' for k, v in style_dict.items()])
                    table['style'] = style_str

            # 2. 处理单元格属性
            for td in soup.find_all('td'):
                attrs_to_style = ['width', 'height', 'align']
                style_dict = {}
                for attr in attrs_to_style:
                    if attr in td.attrs:
                        style_dict[attr] = td[attr]
                        del td[attr]
                if style_dict:
                    existing_style = td.get('style', '')
                    existing_style_dict = self.parse_style(existing_style)
                    style_dict = self.merge_style_dicts(existing_style_dict, style_dict)
                    style_str = ';'.join([f'{k}:{v}' for k, v in style_dict.items()])
                    td['style'] = style_str

            # 3. 处理font标签 - 修改后的版本
            def process_fonts(font_tag):
                """递归处理font标签,将样式合并到最近的非font父标签"""
                parent = font_tag.parent
                if not parent:  # 如果标签已经不在树中，直接返回
                    return None

                # 收集所有font标签的样式
                style_dict = {}
                if 'color' in font_tag.attrs:
                    style_dict['color'] = font_tag['color']
                if 'face' in font_tag.attrs:
                    style_dict['font-family'] = font_tag['face']
                if 'size' in font_tag.attrs:
                    size = font_tag['size']
                    if isinstance(size, str) and 'pt' not in size:
                        # 将数字转换为pt单位
                        style_dict['font-size'] = f'{size}pt'
                    else:
                        style_dict['font-size'] = size
                if 'style' in font_tag.attrs:
                    existing_styles = font_tag['style']
                    existing_style_dict = self.parse_style(existing_styles)
                    style_dict = self.merge_style_dicts(style_dict, existing_style_dict)

                # 处理子元素
                content_tags = []
                text_content = []

                for child in list(font_tag.children):  # 使用list创建副本
                    if isinstance(child, str):
                        if child.strip():
                            text_content.append(child.strip())
                    elif child.name == 'font':
                        # 递归处理嵌套的font标签
                        result = process_fonts(child)
                        if result:
                            if isinstance(result, list):
                                for r in result:
                                    existing_style = r.get('style', '')
                                    r['style'] = self.merge_styles(existing_style, style_dict)
                                    content_tags.append(r)
                            else:
                                existing_style = result.get('style', '')
                                result['style'] = self.merge_styles(existing_style, style_dict)
                                content_tags.append(result)
                    else:
                        # 非font标签，合并样式
                        existing_style = child.get('style', '')
                        child['style'] = self.merge_styles(existing_style, style_dict)
                        content_tags.append(child)

                # 创建新的内容
                if text_content and not content_tags:
                    # 只有文本内容，创建新的span
                    span = soup.new_tag('span')
                    span['style'] = self.merge_styles('', style_dict)
                    span.string = ''.join(text_content)
                    font_tag.replace_with(span)
                    return span
                elif content_tags:
                    font_tag.replace_with(*content_tags)
                    return content_tags
                else:
                    # 没有内容，直接移除
                    font_tag.decompose()
                    return None

            # 从外到内处理font标签
            for font in soup.find_all('font'):
                process_fonts(font)

            # 4. 处理加粗标签
            for b in soup.find_all('b'):
                # 查找父标签中的第一个span标签
                parent = b.parent
                span = parent if parent.name == 'span' else None

                if span:
                    # 如果在父标签中找到span标签，合并样式
                    existing_style = span.get('style', '')
                    style_dict = {}

                    # 解析现有样式
                    if existing_style:
                        existing_style_dict = self.parse_style(existing_style)
                        style_dict = self.merge_style_dicts(style_dict, existing_style_dict)

                    # 添加加粗样式
                    style_dict['font-weight'] = 'bold'

                    b_styles = b.get('style', '')
                    b_style_dict = self.parse_style(b_styles)
                    style_dict = self.merge_style_dicts(style_dict, b_style_dict)

                    # 更新样式
                    span['style'] = ';'.join([f'{k}:{v}' for k, v in style_dict.items()])

                    # 将b标签的内容移动到span中
                    if span.string:
                        if b.get_text() != span.string:
                            span.string = span.string + b.get_text()
                        else:
                            span.string = b.get_text()
                    else:
                        span.string = b.get_text()
                    b.decompose()
                else:
                    # 如果没有找到span标签，创建新的span
                    new_span = soup.new_tag('span')
                    new_span['style'] = 'font-weight:bold'
                    new_span.string = b.get_text()
                    b.replace_with(new_span)

            # 5. 处理span嵌套
            def flatten_spans(span):
                """递归处理span嵌套"""
                # 获取当前span的样式
                parent = span.parent
                if not parent:  # 如果标签已经不在树中，直接返回
                    return None

                outer_style = self.parse_style(span.get('style', ''))

                # 收集直接文本内容
                direct_text = ''
                for child in span.children:
                    if isinstance(child, str) and child.strip():
                        direct_text += child.strip()
                    elif hasattr(child, 'name') and child.name != 'span':
                        # 如果是非span标签，收集其文本内容
                        direct_text += child.get_text().strip()

                # 查找所有直接子span标签
                child_spans = span.find_all('span', recursive=False)

                if not child_spans:
                    # 如果没有子span，返回当前span
                    return span

                if not direct_text:
                    # 如果外层span没有直接文本内容，将样式下沉到子span
                    new_childs = []
                    for child in span.children:
                        child_style = self.parse_style(child.get('style', ''))
                        child_style_dict = self.merge_style_dicts(child_style, outer_style)
                        child['style'] = ';'.join(f'{k}:{v}' for k, v in child_style_dict.items())
                        new_childs.append(child)

                    span.replace_with(*new_childs)
                else:
                    # 如果外层span有直接文本内容，创建新的span包含文
                    new_spans = []
                    # 遍历所有子元素
                    for child in list(span.contents):
                        if isinstance(child, str):
                            if child.strip():
                                # 创建新的span包含文本内容
                                new_span = soup.new_tag('span')
                                new_span['style'] = ';'.join(f'{k}:{v}' for k, v in outer_style.items())
                                new_span.string = child.strip()
                                new_spans.append(new_span)
                        elif child.name == 'span':
                            # 对于子span，合并样式
                            child_style = self.parse_style(child.get('style', ''))
                            child_style = self.merge_style_dicts(child_style, outer_style)
                            child['style'] = ';'.join(f'{k}:{v}' for k, v in child_style.items())
                            new_spans.append(child)
                        else:
                            # 其他类型的标签保持不变
                            new_spans.append(child)
                    span.replace_with(*new_spans)

                # 递归处理所有子span
                for child_span in span.find_all('span', recursive=False):
                    flatten_spans(child_span)

                return span

            # 从外到内处理所有span标签
            for span in soup.find_all('span'):
                flatten_spans(span)
            return str(soup)

        # 先转换成Windows格式
        win_format_html = convert_to_win_format(html_content)
        return self.simple_table_html_win(win_format_html)

    ############################################################
    # docx替换方法
    ############################################################
    def replace_word(self, doc, tag, pv):
        full_text = ''.join(run.text for paragraph in doc.paragraphs for run in paragraph.runs)

        # 如果关键词不在文本中，直接返回
        if tag not in full_text:
            return

        pattern = re.compile(re.escape(tag))
        matches = list(pattern.finditer(full_text))
        # 从后向前替换，以避免索引问题
        for match in reversed(matches):
            start, end = match.span()
            self.replace_text_range(doc, start, end, pv)

    def replace_section_word(self, doc, tag, pv):
        for section in doc.sections:
            self.replace_word(section.header, tag, pv)

    @staticmethod
    def replace_table_word(doc, tag, pv):
        if not isinstance(pv, str):
            return
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if tag in paragraph.text:
                            # 收集所有run的文本和格式
                            runs_info = []
                            for run in paragraph.runs:
                                runs_info.append({
                                    'text': run.text,
                                    'bold': run.bold,
                                    'italic': run.italic,
                                    'underline': run.underline,
                                    'font': run.font.name,
                                    'size': run.font.size,
                                    'color': run.font.color.rgb,
                                    'highlight': run.font.highlight_color
                                })

                            # 合并所有文本并替换标签
                            full_text = ''.join(info['text'] for info in runs_info)
                            new_text = full_text.replace(tag, pv)

                            # 清除段落中的所有运行
                            for run in paragraph.runs:
                                run.clear()

                            # 重新添加文本，保持原有格式
                            new_run = paragraph.add_run(new_text)
                            if runs_info:
                                format_to_apply = runs_info[0]
                                new_run.bold = format_to_apply['bold']
                                new_run.italic = format_to_apply['italic']
                                new_run.underline = format_to_apply['underline']
                                new_run.font.name = format_to_apply['font']
                                new_run.font.size = format_to_apply['size']
                                if format_to_apply['color'] != RGBColor(255, 255, 0):  # 如果不是黄色
                                    new_run.font.color.rgb = format_to_apply['color']
                                if format_to_apply['highlight'] is not None:
                                    new_run.font.highlight_color = None

    @staticmethod
    def replace_text_range(text_frame, start, end, replacement):
        current_index = 0
        start_to_clean = False
        for paragraph in text_frame.paragraphs:
            for run in paragraph.runs:
                run_text = run.text
                run_length = len(run_text)
                run_end = current_index + run_length

                if not start_to_clean and current_index <= start < run_end:
                    # 替换开始的运行
                    prefix = run_text[:start - current_index]
                    run.text = prefix + replacement
                    if end <= run_end:
                        # 替换在同一个运行内结束
                        suffix = run_text[end - current_index:]
                        run.text += suffix
                        return
                    else:
                        start_to_clean = True
                elif start_to_clean and end < run_end:
                    # 替换结束的运行
                    suffix = run_text[end - current_index:]
                    run.text += suffix
                    return
                elif start_to_clean and run_end <= end:
                    run.text = ""
                current_index = run_end

    ############################################################
    # docx查找方法
    ############################################################
    def find_table_by_name(self, doc, doc_html_path, table_name):
        table_html = self.extract_table_from_html(table_name, doc_html_path)
        simple_table_html = self.simple_table_html(table_html)
        return simple_table_html

    def find_paragraph_by_title(self, doc, title):
        content = []
        found = False
        title_level = -1
        for para in doc.paragraphs:
            if para.text.strip().endswith(title):
                found = True
                style = para.style.name.lower()
                match = re.match(r'heading\s*(\d+)', style)
                if match:
                    title_level = int(match.group(1))
                else:
                    title_level = -1
                continue

            if found:
                # 判断当前段落的级别
                current_level = -1
                style = para.style.name.lower()
                match = re.match(r'heading\s*(\d+)', style)
                if match:
                    current_level = int(match.group(1))

                if current_level != -1 and current_level <= title_level:
                    # 遇到同级或更高级的标题，停止收集
                    break

                content.append(para.text.strip())

        return content

    def find_image_by_name(self, doc, image_name, save_dir):
        image_info = ("", None, None)
        try:
            # 获取文档中所有关系
            rels = doc.part.rels

            # 遍历段落查找图片说明
            for para in doc.paragraphs:
                if para.text.startswith('图') and image_name in para.text:
                    # 获取前一个元素（可能包含图片）
                    image_element = para._element.getprevious()
                    if image_element is not None:
                        # 在元素中查找所有图片引用
                        blips = image_element.findall(
                            './/pic:pic//a:blip',
                            {
                                'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
                                'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'
                            }
                        )

                        if blips:
                            # 获取第一个图片的关系ID
                            rId = blips[0].get(qn('r:embed'))
                            if rId and rId in rels:
                                # 获取图片部分
                                image_part = rels[rId].target_part

                                # 创建保存目录
                                if not os.path.exists(save_dir):
                                    os.makedirs(save_dir)

                                # 构建保存路径（使用原始扩展名）
                                image_extension = os.path.splitext(image_part.partname)[1]
                                image_filename = f"{image_name}{image_extension}"
                                image_path = os.path.join(save_dir, image_filename)

                                # 保存图片
                                with open(image_path, 'wb') as f:
                                    f.write(image_part.blob)

                                # 获取图片尺寸信息
                                inline_shapes = image_element.findall(
                                    './/wp:inline',
                                    {'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'}
                                )
                                if inline_shapes:
                                    extent = inline_shapes[0].find('.//wp:extent',
                                                                   {
                                                                       'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'})
                                    if extent is not None:
                                        # 将EMU单位转换为英寸 (1 inch = 914400 EMU)
                                        width = float(extent.get('cx', 0)) / 914400
                                        height = float(extent.get('cy', 0)) / 914400
                                        image_info = (image_path, width, height)
                                        break

                                image_info = (image_path, None, None)
                                break

        except Exception as e:
            print(f"提取图片时发生错误: {str(e)}")
            return ("", None, None)

        return image_info

    def copy_table_to_doc(self, target_doc, target_key, source_doc, doc_html_path, table_name):
        table_html = self.find_table_by_name(source_doc, doc_html_path, table_name)
        html_table_to_docx(target_doc, table_html, target_key)

    def copy_paragraph_to_doc(self, target_doc, target_key, source_doc, paragraph_name):
        paragraph_texts = self.find_paragraph_by_title(source_doc, paragraph_name)
        self.replace_word(target_doc, target_key, '\n\t'.join(paragraph_texts))

    def copy_image_to_doc(self, target_doc, target_key, source_doc, image_name):
        # 创建临时目录
        temp_dir = './temp_images'
        if not os.path.exists(temp_dir):
            os.makedirs(temp_dir)

        # 获取图片
        image_info = self.find_image_by_name(source_doc, image_name, temp_dir)
        image_path = image_info[0]
        if not image_path:
            return
        try:
            # 在目标文档中查找目标段落
            for i, para in enumerate(target_doc.paragraphs):
                if para.text.startswith('图') and para.text.endswith(target_key):
                    # 在目标段落之前插入新段落
                    p = target_doc.add_paragraph()
                    # 将新段落移动到目标段落之前
                    para._p.addprevious(p._p)
                    # 在新段落中添加图片
                    run = p.add_run()
                    run.add_picture(image_path, width=Inches(image_info[1]))
                    break
        finally:
            # 清理临时文件
            if 'image_path' in locals() and os.path.exists(image_path):
                os.remove(image_path)
            if os.path.exists(temp_dir) and not os.listdir(temp_dir):
                os.rmdir(temp_dir)

    ############################################################
    # docx表格插入方法
    ############################################################

    def insert_suolueyu_table(self, doc, key, values):
        table = self.build_table(doc, values)
        self.move_table_to_target(doc, table, key)

    def move_table_to_target(self, doc, table, key):
        target_paragraph = None

        # 遍历所有段落查找目标位置
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip().endswith(key):
                target_paragraph = paragraph
                # 检查下一个元素是否为表格
                next_element = paragraph._element.getnext()
                if next_element.tag.endswith('tbl'):
                    self.remove_element(next_element)
                self.move_table_after(table, target_paragraph)
                break

    @staticmethod
    def remove_element(element):
        """从文档中删除元素"""
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)

    @staticmethod
    def move_table_after(table, paragraph):
        tbl, p = table._tbl, paragraph._p
        p.addnext(tbl)

    def build_table(self, document, result):
        col = 3
        table = document.add_table(rows=0, cols=col)
        no_indent_style = None
        try:
            no_indent_style = document.styles['NoIndent']
        except KeyError:
            # 如果样式不存在,则创建新样式
            no_indent_style = document.styles.add_style('NoIndent', WD_STYLE_TYPE.PARAGRAPH)
            no_indent_style.base_style = document.styles['Normal']
            no_indent_style.paragraph_format.first_line_indent = Pt(0)
            no_indent_style.paragraph_format.left_indent = Pt(0)
            no_indent_style.paragraph_format.space_before = Pt(0)
            no_indent_style.paragraph_format.space_after = Pt(0)
            no_indent_style.paragraph_format.line_spacing = 1.0
        for item in result:
            row_cells = table.add_row().cells
            for i, cell in enumerate(row_cells):
                cell._element.clear_content()
                p = cell.add_paragraph(style='NoIndent')
                pPr = p._element.get_or_add_pPr()
                for child in pPr[:]:
                    pPr.remove(child)
                # 设置缩进XML
                ind = parse_xml(
                    f'<w:ind {nsdecls("w")} w:firstLine="0" w:firstLineChars="0" w:hanging="0" w:left="0" w:right="0"/>')
                pPr.append(ind)

                # 设置段落间距XML
                spacing = parse_xml(
                    f'<w:spacing {nsdecls("w")} w:before="0" w:after="0" w:line="240" w:lineRule="auto"/>')
                pPr.append(spacing)
                run = p.add_run(item[i])
                font = run.font
                if self.is_english(item[i]):
                    font.name = 'Times New Roman'
                else:
                    font.name = 'SimSun'
        return table

    @staticmethod
    def is_english(text):
        """
        判断文本是否为英文

        Args:
            text (str): 需要判断的文本

        Returns:
            bool: 如果文本全部为英文字符返回True，否则返回False
        """
        # 移除所有空白字符和标点符号
        text = re.sub(r'[\s\.,!?;:\'\"()\[\]{}]', '', text)
        if not text:  # 如果移除后为空，返回False
            return False
        # 检查剩余字符是否全部为英文字母
        return bool(re.match(r'^[a-zA-Z]+$', text))

    @staticmethod
    def is_chinese(text):
        """
        判断文本是否为中文

        Args:
            text (str): 需要判断的文本

        Returns:
            bool: 如果文本包含中文字符返回True，否则返回False
        """
        # \u4e00-\u9fff 是中文字符的Unicode范围
        return bool(re.search('[\u4e00-\u9fff]', text))
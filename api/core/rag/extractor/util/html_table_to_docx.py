from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
import re


def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    for edge in ['top', 'left', 'bottom', 'right']:
        if edge in kwargs:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            element.set(qn('w:val'), 'single')
            element.set(qn('w:sz'), kwargs[edge])
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), '000000')
            tcPr.append(element)


def set_cell_background(cell, color):
    """设置单元格背景色"""
    tcPr = cell._tc.get_or_add_tcPr()
    shading_elm = parse_xml(
        f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{color}"/>')
    tcPr.append(shading_elm)


def parse_color(color_str):
    """解析颜色值"""
    if color_str.startswith('#'):
        return color_str[1:]
    elif color_str.lower() == 'a9a9a9':  # 处理特定的灰色
        return 'A9A9A9'
    return None


def get_style_dict(cell):
    """从HTML单元格中提取样式信息"""
    style_dict = {}

    # 获取style属性
    style_attr = cell.get('style', '')
    if style_attr:
        # 按分号分割样式
        styles = [s.strip() for s in style_attr.split(';') if s.strip()]

        for style in styles:
            # 处理width
            if style.startswith('width:'):
                style_dict['width'] = style.replace('width:', '').strip()

            # 处理background-color
            elif style.startswith('background-color:'):
                style_dict['background-color'] = style.replace('background-color:', '').strip()

            # 处理text-align
            elif style.startswith('text-align:'):
                style_dict['text-align'] = style.replace('text-align:', '').strip()

            elif style.startswith('vertical-align:'):
                style_dict['vertical-align'] = style.replace('vertical-align:', '').strip()

    # 检查p标签的样式
    for p in cell.find_all('p'):
        p_style = p.get('style', '')
        if p_style:
            styles = [s.strip() for s in p_style.split(';') if s.strip()]
            for style in styles:
                if style.startswith('text-align:'):
                    style_dict['text-align'] = style.replace('text-align:', '').strip()

    # 检查span标签中的样式
    for span in cell.find_all('span'):
        span_style = span.get('style', '')
        if span_style:
            styles = [s.strip() for s in span_style.split(';') if s.strip()]
            for style in styles:
                if style.startswith('background-color:'):
                    style_dict['background-color'] = style.replace('background-color:', '').strip()
                elif style.startswith('font-weight:'):
                    style_dict['font-weight'] = style.replace('font-weight:', '').strip()

    return style_dict


def get_alignment(align_str):
    """转换对齐方式"""
    align_map = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
    }
    return align_map.get(align_str.lower(), WD_ALIGN_PARAGRAPH.LEFT)


def get_vertical_alignment(align_str):
    """转换垂直对齐方式"""
    align_map = {
        'top': WD_ALIGN_VERTICAL.TOP,
        'middle': WD_ALIGN_VERTICAL.CENTER,
        'bottom': WD_ALIGN_VERTICAL.BOTTOM
    }
    return align_map.get(align_str.lower(), WD_ALIGN_VERTICAL.TOP)


def parse_indent(indent_str):
    """解析缩进值"""
    if not indent_str:
        return None

    # 处理pt单位
    if indent_str.endswith('pt'):
        return Pt(float(indent_str.rstrip('pt')))

    # 处理其他单位的缩进值（如果需要）
    # ...

    return None


def format_text(text):
    return text.replace('\n', ' ').replace('\r', ' ').replace('\t', ' ').replace('  ', ' ').strip()


def get_cell_content(cell):
    """提取单元格内容和样式"""
    paragraphs = []

    # 首先尝试查找p标签
    p_tags = cell.find_all('p')
    if p_tags:
        for p in p_tags:
            text = ""
            # 获取段落样式
            style_dict = {}
            p_style = p.get('style', '')
            if p_style:
                styles = [s.strip() for s in p_style.split(';') if s.strip()]
                for style in styles:
                    if style.startswith('text-align:'):
                        style_dict['text-align'] = style.replace('text-align:', '').strip()
                    elif style.startswith('text-indent:'):
                        style_dict['text-indent'] = True
                    elif style.startswith('line-height:'):
                        style_dict['line-height'] = style.replace('line-height:', '').strip()

            # 获取文本内容，保留制表符
            spans = p.find_all('span')

            span_data = []
            for span in spans:
                span_style_dict = {}
                span_style = span.get('style', '')
                span_text = span.get_text()
                if span_style:
                    s_styles = [s.strip() for s in span_style.split(';') if s.strip()]
                    for s_style in s_styles:
                        if s_style.startswith('vertical-align:'):
                            span_style_dict['vertical-align'] = s_style.replace('vertical-align:', '').strip()
                        elif s_style.startswith('background-color:'):
                            span_style_dict['background-color'] = parse_color(
                                s_style.replace('background-color:', '').strip())
                        elif s_style.startswith('font-weight:'):
                            span_style_dict['font-weight'] = s_style.replace('font-weight:', '').strip()
                span_data.append({
                    'text': format_text(span_text),
                    'style': span_style_dict
                })
            if not spans:
                text = p.get_text()  # 不使用strip=True
            paragraphs.append({
                'text': format_text(text),  # 只在最后去除首尾空白
                'style': style_dict,
                'span_data': span_data
            })
    else:
        # 如果没有p标签，将整个内容作为一个段落
        text = cell.get_text()  # 不使用strip=True
        style_dict = {}

        paragraphs.append({
            'text': format_text(text),  # 只在最后去除首尾空白
            'style': style_dict
        })

    return paragraphs


def get_line_height(style_dict):
    """解析行高值"""
    if 'line-height' in style_dict:
        line_height = style_dict['line-height']
        # 处理像素值
        if line_height.endswith('px'):
            pixels = float(line_height.rstrip('px'))
            return Pt(pixels)
        # 处理数字（倍数）
        elif line_height.replace('.', '').isdigit():
            return float(line_height)
        # 处理百分比
        elif line_height.endswith('%'):
            percentage = float(line_height.rstrip('%'))
            return percentage / 100
    return None


def format_cell(doc_cell, cell, paragraphs_info, is_header=False):
    """设置单元格格式"""
    # 清除现有内容
    doc_cell._element.clear_content()

    # 获取单元格级别的样式
    cell_style = get_style_dict(cell)

    if 'vertical-align' in cell_style:
        doc_cell.vertical_alignment = get_vertical_alignment(cell_style['vertical-align'])

    # 设置单元格边框
    set_cell_border(
        doc_cell,
        top="4",
        bottom="4",
        left="4",
        right="4",
    )

    # 设置背景色
    if 'background-color' in cell_style:
        color = parse_color(cell_style['background-color'])
        if color:
            set_cell_background(doc_cell, color)

    # 确保paragraphs_info是列表
    if isinstance(paragraphs_info, str):
        paragraphs_info = [{'text': paragraphs_info, 'style': {}}]
    elif not isinstance(paragraphs_info, list):
        paragraphs_info = [{'text': str(paragraphs_info), 'style': {}}]

    # 为每个段落设置格式
    for para_info in paragraphs_info:
        paragraph = doc_cell.add_paragraph()

        # 获取段落样式
        para_style = para_info.get('style', {})

        # 合并单元格和段落的样式
        style_dict = dict(cell_style)
        style_dict.update(para_style)

        # 设置段后间距为0
        paragraph.paragraph_format.space_after = Pt(0)

        line_height = get_line_height(style_dict)
        if line_height:
            if isinstance(line_height, float):
                # 如果是倍数，使用multiple
                paragraph.paragraph_format.line_spacing = line_height
            else:
                # 如果是固定值（Pt），使用exactly
                paragraph.paragraph_format.line_spacing = line_height
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY

        # 设置对齐方式
        if 'text-align' in style_dict:
            paragraph.alignment = get_alignment(style_dict['text-align'])
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 添加文本并设置格式
        if len(para_info.get('span_data', [])) == 0:
            text = para_info.get('text', '')
            run = paragraph.add_run(text)
            font = run.font
            font.size = Pt(10.5)
            font.name = 'Times New Roman'
            if 'font-weight' in style_dict:
                run.bold = True if style_dict['font-weight'] == 'bold' else False
        else:
            for span in para_info['span_data']:
                run = paragraph.add_run(span['text'])
                font = run.font
                font.size = Pt(10.5)
                font.name = 'Times New Roman'
                if 'font-weight' in span['style']:
                    run.bold = True if span['style']['font-weight'] == 'bold' else False
                if 'vertical-align' in span['style']:
                    vertical_align = span['style']['vertical-align']
                    if vertical_align == 'sub':
                        run.font.subscript = True
                    elif vertical_align == 'super':
                        run.font.superscript = True

        # 如果需要制表符缩进
        if style_dict.get('text-indent', False):
            paragraph.paragraph_format.left_indent = Pt(21)


def get_merged_cells_info(table):
    """获取合并单元格信息"""
    merged_cells = []
    for i, row in enumerate(table.find_all('tr')):
        for j, cell in enumerate(row.find_all(['td', 'th'])):
            rowspan = int(cell.get('rowspan', 1))
            colspan = int(cell.get('colspan', 1))
            if rowspan > 1 or colspan > 1:
                merged_cells.append({
                    'row': i,
                    'col': j,
                    'rowspan': rowspan,
                    'colspan': colspan
                })
    return merged_cells


def get_column_widths(table):
    """获取表格的列宽信息"""
    widths = []
    total_pixels = 0
    first_row = table.find('tr')
    if first_row:
        cells = first_row.find_all(['td', 'th'])
        for cell in cells:
            style_dict = get_style_dict(cell)
            width = None
            if 'width' in style_dict:
                width_str = style_dict['width']
                # 如果是像素值
                if width_str.endswith('px'):
                    width_pixels = float(width_str.rstrip('px'))
                    total_pixels += width_pixels
                    width = width_pixels  # 暂时保存像素值
                if width_str.endswith('%'):
                    width = float(width_str.rstrip('%'))
            widths.append(width)

    # 将像素值转换为相对比例
    if total_pixels > 0:
        widths = [float(w * 100 / total_pixels) if w else None for w in widths]

    return widths


def get_table_style_dict(table):
    """从HTML表格中提取样式信息"""
    style_dict = {}

    # 获取table的style属性
    style_attr = table.get('style', '')
    if style_attr:
        # 按分号分割样式
        styles = [s.strip() for s in style_attr.split(';') if s.strip()]

        for style in styles:
            # 处理table-layout
            if style.startswith('table-layout:'):
                style_dict['table-layout'] = style.replace('table-layout:', '').strip()

            # 处理width
            elif style.startswith('width:'):
                style_dict['width'] = style.replace('width:', '').strip()

            elif style.startswith('total_width:'):
                style_dict['total_width'] = style.replace('total_width:', '').strip()

            # 处理margin
            elif style.startswith('margin-'):
                style_dict[style.split(':')[0].strip()] = style.split(':')[1].strip()

            # 处理border-collapse
            elif style.startswith('border-collapse:'):
                style_dict['border-collapse'] = style.replace('border-collapse:', '').strip()

    return style_dict


def set_column_width(column, width_percent):
    """设置列宽"""
    for cell in column.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        pct_width = int(width_percent * 50)
        tcW.set(qn('w:w'), str(pct_width))
        tcW.set(qn('w:type'), 'pct')
        tcPr.append(tcW)


def get_row_height(row):
    """获取行高"""
    style_attr = row.get('style', '')
    if style_attr:
        styles = [s.strip() for s in style_attr.split(';') if s.strip()]
        for style in styles:
            if style.startswith('height:'):
                height_str = style.replace('height:', '').strip()
                # 处理像素值
                if height_str.endswith('px'):
                    height_pixels = float(height_str.rstrip('px'))
                    return Pt(height_pixels)  # 将px转换为pt
    return None


def set_row_height(row, height):
    """设置行高"""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(height.twips))
    trHeight.set(qn('w:hRule'), 'atLeast')  # 或者使用 'atLeast' 作为最小高度
    trPr.append(trHeight)


def mv_table_to_target(doc, doc_table, key):
    target_paragraph = None

    # 遍历所有段落查找目标位置
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().endswith(key):
            target_paragraph = paragraph
            # 检查下一个元素是否为表格
            next_element = paragraph._element.getnext()
            if next_element.tag.endswith('tbl'):
                remove_element(next_element)
            move_table_after(doc_table, target_paragraph)
            break


def remove_element(element):
    """从文档中删除元素"""
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def html_table_to_docx(doc, html, key):
    # 解析HTML
    soup = BeautifulSoup(html, 'html.parser')

    # 查找表格
    table = soup.find('table')
    if not table:
        raise ValueError("未找到表格")

    # 获取表格样式信息
    table_style = get_table_style_dict(table)

    # 获取所有行
    rows = table.find_all('tr')
    if not rows:
        raise ValueError("表格中未找到行")

    # 获取列宽信息
    column_widths = get_column_widths(table)

    # 计算最大列数
    max_columns = max(
        sum(int(cell.get('colspan', 1)) for cell in row.find_all(['td', 'th']))
        for row in rows
    )

    # 创建Word表格
    doc_table = doc.add_table(rows=len(rows), cols=max_columns, style='Table Grid')

    # 设置列宽

    if column_widths:
        # 调整列宽
        for i, width in enumerate(column_widths):
            if width:
                set_column_width(doc_table.columns[i], width)

    # 用于跟踪已合并的单元格
    skip_cells = set()

    # 填充表格
    for i, row in enumerate(rows):
        row_height = get_row_height(row)
        if row_height:
            set_row_height(doc_table.rows[i], row_height)

        col_idx = 0
        for cell in row.find_all(['td', 'th']):
            # 跳过已合并的单元格
            while (i, col_idx) in skip_cells:
                col_idx += 1

            # 获取rowspan和colspan
            rowspan = int(cell.get('rowspan', 1))
            colspan = int(cell.get('colspan', 1))

            # 设置单元格内容和格式
            if colspan > 1 or rowspan > 1:
                # 合并单元格
                target_cell = doc_table.cell(i, col_idx)
                if colspan > 1:
                    target_cell.merge(doc_table.cell(i, col_idx + colspan - 1))
                if rowspan > 1:
                    target_cell.merge(doc_table.cell(i + rowspan - 1, col_idx))

                # 标记被合并的单元格
                for r in range(i, i + rowspan):
                    for c in range(col_idx, col_idx + colspan):
                        if (r, c) != (i, col_idx):
                            skip_cells.add((r, c))

            # 设置内容和格式
            if (i, col_idx) not in skip_cells:
                paragraphs_info = get_cell_content(cell)
                format_cell(doc_table.cell(i, col_idx), cell, paragraphs_info, i == 0)

            col_idx += colspan

    mv_table_to_target(doc, doc_table, key)


if __name__ == "__main__":
    # 示例用法
    with open('test.html', 'r', encoding='utf-8') as f:
        html_content = f.read()
    doc = Document('output.docx')
    html_table_to_docx(doc, html_content, '试验及药品基本信息')
    doc.save('output.docx')